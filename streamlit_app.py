import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
import io

# ==========================================
# --- APP CONFIGURATION ---
# ==========================================
st.set_page_config(page_title="Reaction Forces Viewer", layout="wide")
st.title("🏗️ Structural Reaction Forces Viewer")
st.markdown("Upload your `.docx` output file to explore reaction extremes and load combinations.")

# ==========================================
# --- HELPER FUNCTIONS ---
# ==========================================
def get_extreme_color_scale(component_type):
    return 'YlGnBu_r' if "MIN" in component_type else 'YlOrRd'

def get_lc_color_scale(values):
    v_min, v_max = values.min(), values.max()
    if v_min < 0 and v_max > 0:
        return 'rainbow'
    return 'YlOrRd' if abs(v_max) > abs(v_min) else 'YlGnBu_r'

def adjust_figure_size(fig, df_plot, padding=10.0, scale=1.0):
    """Calculates optimal width/height with a user-defined scale factor."""
    x_coords = pd.to_numeric(df_plot['X [m]'], errors='coerce').dropna()
    y_coords = pd.to_numeric(df_plot['Y [m]'], errors='coerce').dropna()

    if x_coords.empty or y_coords.empty:
        fig.update_layout(width=int(800 * scale), height=int(600 * scale))
        return

    x_min, x_max = x_coords.min(), x_coords.max()
    y_min, y_max = y_coords.min(), y_coords.max()
    
    x_span = max((x_max - x_min) + (2 * padding), 1.0)
    y_span = max((y_max - y_min) + (2 * padding), 1.0)
    
    # Increase the base size and apply the scale factor
    base_size = 800 * scale 
    
    if x_span >= y_span:
        width = base_size
        # The +120 is for the legend/title; we scale a portion of that too
        height = (y_span / x_span) * base_size + (100 * scale) 
    else:
        height = base_size
        width = (x_span / y_span) * base_size + (150 * scale) 
        
    fig.update_layout(width=int(width), height=int(height))
    fig.update_xaxes(range=[x_min - padding, x_max + padding])
    fig.update_yaxes(range=[y_min - padding, y_max + padding])

def get_table_as_df(table):
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = [cell.text.strip() for cell in row.cells]
        if i == 0:
            keys = text
            continue
        data.append(dict(zip(keys, text)))
    return pd.DataFrame(data)

@st.cache_data
def extract_raw_tables(file_bytes):
    """Only extracts data. No renaming, no cleaning."""
    doc = Document(io.BytesIO(file_bytes))
    raw_data = []
    for i, table in enumerate(doc.tables):
        df = get_table_as_df(table)
        raw_data.append(df)
    return raw_data

# ==========================================
# --- DATA PROCESSING (CACHED) ---
# ==========================================
@st.cache_data
def load_and_process_data(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    
    # --- 1. EXTRACTION ---
    # Using indices directly since we confirmed they are 0, 1, 2
    df_points_raw = get_table_as_df(doc.tables[0])
    df_supports_raw = get_table_as_df(doc.tables[1])
    df_loadcombinations_raw = get_table_as_df(doc.tables[2])

    # --- 2. FORCED RENAMING (Aligned to your Debug Info) ---
    # Table 0: ['Ordinal', 'X [m]', 'Y [m]', 'Z [m]']
    df_points_raw = df_points_raw.iloc[:, :4]
    df_points_raw.columns = ['Ordinal', 'X [m]', 'Y [m]', 'Z [m]']
    
    # Table 1: ['Name', 'Supported point (number)', 'Type', 'Support coordinate system', '']
    # We take the first 4 and map them correctly
    df_supports_raw = df_supports_raw.iloc[:, :4]
    df_supports_raw.columns = ['Name', 'PointNum', 'Type', 'Sys']
    
    # Table 2: ['Support name', 'Load combinations', 'Fx [kN]', ...]
    df_loadcombinations_raw = df_loadcombinations_raw.iloc[:, :8]
    df_loadcombinations_raw.columns = ['SupName', 'LC', 'Fx [kN]', 'Fy [kN]', 'Fz [kN]', 'Mx [kNm]', 'My [kNm]', 'Mz [kNm]']

    # --- 3. CLEANING & NORMALIZATION ---
    # Points: Convert coords to float
    df_points = df_points_raw.copy()
    for col in ['X [m]', 'Y [m]', 'Z [m]']:
        df_points[col] = pd.to_numeric(df_points[col].astype(str).str.replace(',', '.'), errors='coerce')
    # Force Ordinal to string for clean merging
    df_points['Ordinal'] = df_points['Ordinal'].astype(str).str.strip()

    # Supports: Clean Name and PointNum
    df_supports = df_supports_raw.copy()
    df_supports['Name'] = df_supports['Name'].astype(str).str.strip()
    df_supports['PointNum'] = df_supports['PointNum'].astype(str).str.strip()
    
    # Reactions: ffill, strip names, and convert to numeric
    df_loadcombinations = df_loadcombinations_raw.copy()
    df_loadcombinations['SupName'] = df_loadcombinations['SupName'].replace('', pd.NA).ffill().astype(str).str.strip()
    
    reaction_cols = ['Fx [kN]', 'Fy [kN]', 'Fz [kN]', 'Mx [kNm]', 'My [kNm]', 'Mz [kNm]']
    for col in reaction_cols:
        df_loadcombinations[col] = pd.to_numeric(df_loadcombinations[col].astype(str).str.replace(',', '.'), errors='coerce')

    # --- 4. MERGING ---
    # Step A: Link Supports to Coordinates
    df_support_coords = pd.merge(
        df_supports, 
        df_points[['Ordinal', 'X [m]', 'Y [m]']], 
        left_on='PointNum', 
        right_on='Ordinal', 
        how='left'
    )
    
    # Step B: Link Reactions to the Coordinate-enabled Supports
    df_final = pd.merge(
        df_loadcombinations, 
        df_support_coords[['Name', 'X [m]', 'Y [m]']], 
        left_on='SupName', right_on='Name', how='left'
    )
    
    # Standardize column names and drop redundant 'Name' column from the merge
    df_final = df_final.rename(columns={'SupName': 'Support name', 'LC': 'Load combinations'})
    df_final = df_final.drop(columns=['Name'], errors='ignore') 
    df_final = df_final.dropna(subset=['X [m]', 'Y [m]'])

    if df_final.empty:
        return None, None, "Merge Error: Could not match Support Names to Coordinates. Check Table 1 vs Table 2."

    # --- 5. CALCULATE EXTREMES ---
    extreme_list = []
    for support_name, group in df_final.groupby('Support name'):
        for col in reaction_cols:
            clean_group = group.dropna(subset=[col])
            if clean_group.empty: continue
            
            # Min
            idx_min = clean_group[col].idxmin()
            row_min = clean_group.loc[[idx_min]].copy()
            row_min['Component Type'] = f"{col.split(' ')[0]}_MIN"
            extreme_list.append(row_min)
            
            # Max
            idx_max = clean_group[col].idxmax()
            row_max = clean_group.loc[[idx_max]].copy()
            row_max['Component Type'] = f"{col.split(' ')[0]}_MAX"
            extreme_list.append(row_max)

    # After the loop finishes:
    df_reaction_extremes = pd.concat(extreme_list).reset_index(drop=True)

    # Drop the redundant 'Name' column from extremes
    df_reaction_extremes = df_reaction_extremes.drop(columns=['Name'], errors='ignore')

    # Reorder columns: Move 'Component Type' to index 1 (second place)
    cols = list(df_reaction_extremes.columns)
    if 'Component Type' in cols:
        # Remove it from its current position and insert it at index 1
        cols.insert(1, cols.pop(cols.index('Component Type')))
    
    df_reaction_extremes = df_reaction_extremes[cols]

    return df_final, df_reaction_extremes, None

# ==========================================
# --- PLOTTING LOGIC ---
# ==========================================
def plot_extremes(df, component_type, scale=1.0, marker_size=12, text_size=10):
    base = component_type.split('_')[0]
    unit = "[kN]" if base.startswith('F') else "[kNm]"
    target_col = f"{base} {unit}"
    
    df_plot = df[df['Component Type'] == component_type].copy()
    if df_plot.empty: return None
    
    # Calculate values for the title
    v_min = df_plot[target_col].min()
    v_max = df_plot[target_col].max()

    fig = px.scatter(
        df_plot, x='X [m]', y='Y [m]', color=target_col, text=target_col,
        color_continuous_scale=get_extreme_color_scale(component_type),
        title=f"Reaction Extremes: {component_type} <{v_min:.1f}; {v_max:.1f}>",
        labels={target_col: f"Value {unit}"},
        hover_data=['Support name', 'Load combinations']
    )
    
    fig.update_traces(
        mode='markers+text', 
        texttemplate='%{text:.0f}', 
        textfont=dict(size=text_size),  # <--- FIX: use the variable
        marker=dict(size=marker_size, line=dict(width=1.5, color='black')), # <--- FIX: use the variable
        textposition='top right'
    )
    
    adjust_figure_size(fig, df_plot, padding=10.0, scale=scale)
    return fig

def plot_load_combination(df, load_comb, force_component, scale=1.0, marker_size=12, text_size=10):
    """
    Generates a scatter plot for a specific Load Combination using 
    centralized color logic and user-defined sizing.
    """
    # 1. Filter data
    df_plot = df[df['Load combinations'] == load_comb].copy()
    if df_plot.empty: 
        return None

    # 2. Get the colorscale from your helper method
    color_scheme = get_lc_color_scale(df_plot[force_component])
    
    # Determine if we need a midpoint (useful for diverging scales like 'rainbow' or 'RdBu')
    v_min = df_plot[force_component].min()
    v_max = df_plot[force_component].max()
    c_mid = 0 if (v_min < 0 and v_max > 0) else None

    # 3. Create Figure
    fig = px.scatter(
        df_plot, 
        x='X [m]', 
        y='Y [m]', 
        color=force_component, 
        text=force_component,
        color_continuous_scale=color_scheme,
        color_continuous_midpoint=c_mid,
        title=f"LC: {load_comb} | {force_component} <{v_min:.1f}; {v_max:.1f}>",
        labels={force_component: "Value"},
        hover_data=['Support name']
    )
    
    # 4. Apply UI controls (Marker & Text Size)
    fig.update_traces(
        mode='markers+text', 
        texttemplate='%{text:.0f}', 
        textfont=dict(size=text_size), 
        marker=dict(
            size=marker_size, 
            line=dict(width=1.5, color='black')
        ), 
        textposition='top right'
    )
    
    # 5. Maintain structural aspect ratio and global scale
    fig.update_yaxes(scaleanchor="x", scaleratio=1)
    adjust_figure_size(fig, df_plot, padding=10.0, scale=scale)
    
    return fig

# ==========================================
# --- STREAMLIT USER INTERFACE ---
# ==========================================

uploaded_file = st.file_uploader("Upload Word Document (.docx)", type=['docx'])

if uploaded_file:
    file_bytes = uploaded_file.read()
    
    # # DEBUG SECTION: Show everything we found in the Word doc
    # raw_tables = extract_raw_tables(file_bytes)
    # with st.expander("🛠️ DEBUG: Raw Imported Tables (Inspect these first!)", expanded=True):
    #     st.write(f"Total tables found in document: {len(raw_tables)}")
    #     for i, table_df in enumerate(raw_tables):
    #         st.markdown(f"**Table {i}** (Columns: {list(table_df.columns)})")
    #         st.dataframe(table_df.head(10))

    # Try to process the data
    df_final, df_reaction_extremes, error_msg = load_and_process_data(file_bytes)

    if error_msg:
        st.error(f"Processing Error: {error_msg}")
    else:
        st.success("Data processed successfully!")
    
    # --- GLOBAL PLOT CONTROLS ---
    st.markdown("### 🛠️ Plot Appearance")
    ctrl_col1, ctrl_col2, ctrl_col3 = st.columns(3)

    with ctrl_col1:
        plot_scale = st.slider("📏 Plot Scale", 0.5, 3.0, 1.0, 0.1)
    with ctrl_col2:
        marker_size = st.slider("🔵 Marker Size", 5, 40, 12, 1)
    with ctrl_col3:
        text_size = st.slider("🔤 Text Size", 6, 24, 10, 1)

    # 2. Create Tabs for the UI
    tab1, tab2, tab3 = st.tabs(["🔥 Reaction Extremes", "📊 Load Combinations", "📋 Raw Data"])

    # --- TAB 1: EXTREMES ---
    with tab1:
        st.subheader("Global Component Extremes")
        comp_types = df_reaction_extremes['Component Type'].unique()
        selected_extreme = st.selectbox("Select Component Extreme:", sorted(comp_types))
        
        if selected_extreme:
            # CRITICAL: You must pass marker_size and text_size here!
            fig_ex = plot_extremes(
                df_reaction_extremes, 
                selected_extreme, 
                scale=plot_scale, 
                marker_size=marker_size, # Added this
                text_size=text_size      # Added this
            )
            if fig_ex:
                st.plotly_chart(fig_ex, width="content")

    # --- TAB 2: LOAD COMBINATIONS ---
    with tab2:
        st.subheader("Specific Load Combinations")
        col_lc1, col_lc2 = st.columns(2)
        
        with col_lc1:
            unique_lcs = df_final['Load combinations'].dropna().unique()
            selected_lc = st.selectbox("Select Load Combination:", sorted(unique_lcs))
        with col_lc2:
            force_components = ['Fx [kN]', 'Fy [kN]', 'Fz [kN]', 'Mx [kNm]', 'My [kNm]', 'Mz [kNm]']
            selected_force = st.selectbox("Select Force Component:", force_components)
                
        if selected_lc and selected_force:
            # Generate Plot with all controls
            fig_lc = plot_load_combination(df_final, selected_lc, selected_force, 
                                        scale=plot_scale, marker_size=marker_size, text_size=text_size)
            if fig_lc:
                st.plotly_chart(fig_lc, width="content")
            
            # --- DOWNLOAD LC DATA ---
            st.divider()
            df_lc_filtered = df_final[df_final['Load combinations'] == selected_lc].copy()
            
            st.write(f"📥 Download data for **{selected_lc}**")
            base_lc_name = uploaded_file.name.rsplit('.', 1)[0]
            lc_filename = f"{base_lc_name}_{selected_lc}"

            dl_col1, dl_col2 = st.columns(2)
            with dl_col1:
                st.download_button(
                    label=f"Download {selected_lc} as CSV",
                    data=df_lc_filtered.to_csv(index=False).encode('utf-8'),
                    file_name=f"{lc_filename}.csv",
                    mime='text/csv',
                    use_container_width=True
                )
            with dl_col2:
                buffer_lc = io.BytesIO()
                with pd.ExcelWriter(buffer_lc, engine='xlsxwriter') as writer:
                    df_lc_filtered.to_excel(writer, index=False, sheet_name=selected_lc[:31]) # Excel sheet name limit
                st.download_button(
                    label=f"Download {selected_lc} as XLSX",
                    data=buffer_lc.getvalue(),
                    file_name=f"{lc_filename}.xlsx",
                    mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    use_container_width=True
                )
                
    # --- TAB 3: RAW DATA (Bonus) ---
    with tab3:
        st.subheader("Processed DataFrame (`df_final`)")
        st.dataframe(df_final, width='stretch')
        st.subheader("Extremes DataFrame (`df_reaction_extremes`)")
        st.dataframe(df_reaction_extremes, width='stretch')

        # --- EXPORT SECTION ---
        if uploaded_file and df_reaction_extremes is not None:
            st.divider()
            st.subheader("📥 Export Extreme Results")
            
            # 1. Prepare Filename
            # uploaded_file.name gives something like "MyProject.docx"
            base_name = uploaded_file.name.rsplit('.', 1)[0]
            xlsx_filename = f"{base_name}_Reaction Extremes.xlsx"
            csv_filename = f"{base_name}_Reaction Extremes.csv"

            col_dl1, col_dl2 = st.columns(2)

            with col_dl1:
                # --- CSV Export ---
                csv_data = df_reaction_extremes.to_csv(index=False).encode('utf-8')
                st.download_button(
                    label="Download as CSV",
                    data=csv_data,
                    file_name=csv_filename,
                    mime='text/csv',
                    use_container_width=True
                )

            with col_dl2:
                # --- Excel Export ---
                buffer = io.BytesIO()
                with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
                    df_reaction_extremes.to_excel(writer, index=False, sheet_name='Extremes')
                
                st.download_button(
                    label="Download as XLSX",
                    data=buffer.getvalue(),
                    file_name=xlsx_filename,
                    mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    use_container_width=True
                )